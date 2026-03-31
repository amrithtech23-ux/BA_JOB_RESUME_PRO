import streamlit as st
import requests
import os
import tempfile
import re
from datetime import datetime
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import inch
from reportlab.lib.colors import black, gray

# ============================================
# Page Configuration
# ============================================
st.set_page_config(
    page_title="Business Analyst Job Apply Pro",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============================================
# Custom CSS
# ============================================
st.markdown("""
    <style>
    .report-box {
        font-family: 'Courier New', Courier, monospace;
        white-space: pre;
        background-color: #0e1117;
        color: #00ff00;
        padding: 20px;
        border-radius: 10px;
        overflow-x: auto;
        font-size: 11px;
    }
    .success-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        margin: 1rem 0;
    }
    .warning-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        margin: 1rem 0;
    }
    .error-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
        margin: 1rem 0;
    }
    .stButton>button {
        width: 100%;
    }
    [data-testid="stSidebar"] {
        display: none;
    }
    </style>
""", unsafe_allow_html=True)

# ============================================
# Initialize Session State
# ============================================
if 'reset_counter' not in st.session_state:
    st.session_state.reset_counter = 0
if 'validation_report' not in st.session_state:
    st.session_state.validation_report = None
if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'processing' not in st.session_state:
    st.session_state.processing = False

# ============================================
# Helper Functions
# ============================================

def clean_text_for_api(text):
    """Clean text before sending to API - remove problematic characters."""
    if not text:
        return ""
    
    # Remove or replace problematic characters
    text = text.replace('\x00', '')  # Null characters
    text = text.replace('\x1b', '')  # Escape characters
    text = text.replace('\f', '')    # Form feed
    text = text.replace('\v', '')    # Vertical tab
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove control characters
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
    
    return text.strip()

def extract_text_from_file(uploaded_file):
    """Extracts text from PDF, DOCX, or TXT files with better error handling."""
    try:
        text = ""
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif file_extension == 'docx':
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
        elif file_extension == 'txt':
            text = uploaded_file.read().decode("utf-8", errors='ignore')
        else:
            st.error(f"❌ Unsupported file type: .{file_extension}")
            return None
        
        # Clean the extracted text
        text = clean_text_for_api(text)
        return text.strip()
    
    except Exception as e:
        st.error(f"❌ Error reading file: {str(e)}")
        return None

def get_api_key():
    """Get API key from secrets or user input."""
    try:
        if hasattr(st, 'secrets') and "OPENROUTER_API_KEY" in st.secrets:
            return st.secrets["OPENROUTER_API_KEY"]
    except:
        pass
    return st.text_input("OpenRouter API Key", type="password", 
                        help="Get your key from https://openrouter.ai/keys",
                        key=f"api_key_input_{st.session_state.reset_counter}")

def test_api_connection(api_key):
    """Test if API key is valid."""
    try:
        url = "https://openrouter.ai/api/v1/auth/key"
        headers = {"Authorization": f"Bearer {api_key}"}
        response = requests.get(url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            if 'data' in data:
                st.success(f"✅ API Key Valid! Email: {data['data'].get('email', 'Unknown')}")
                return True
        else:
            st.error(f"❌ API Key Invalid: {response.status_code}")
            return False
    except Exception as e:
        st.error(f"❌ Connection test failed: {str(e)}")
        return False

def call_openrouter_api(prompt, system_instruction, api_key, model="qwen/qwen-2.5-72b-instruct", max_tokens=4000):
    """Calls OpenRouter API with comprehensive error handling."""
    
    # Validate inputs
    if not prompt or not prompt.strip():
        st.error("❌ Error: Prompt is empty")
        return None
    
    if not system_instruction or not system_instruction.strip():
        st.error("❌ Error: System instruction is empty")
        return None
    
    # Clean the prompt
    prompt = clean_text_for_api(prompt)
    system_instruction = clean_text_for_api(system_instruction)
    
    # Check prompt length (OpenRouter has limits)
    prompt_length = len(prompt) + len(system_instruction)
    max_safe_length = 60000  # Conservative limit for safety
    
    if prompt_length > max_safe_length:
        st.warning(f"⚠️ Prompt is very long ({prompt_length} chars). Truncating...")
        # Truncate prompt while keeping important parts
        max_prompt_length = max_safe_length - len(system_instruction) - 1000
        prompt = prompt[:max_prompt_length] + "\n\n[Content truncated due to length]"
    
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/amrithtech23-ux/BA_JOB_RESUME_PRO",
        "X-Title": "BA Job Apply Pro"
    }
    
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens
    }
    
    try:
        with st.spinner("🔄 Sending request to API..."):
            response = requests.post(url, headers=headers, json=payload, timeout=120)
        
        # Detailed error handling
        if response.status_code != 200:
            error_text = response.text[:500] if response.text else "No error details"
            
            st.error(f"❌ API Error ({response.status_code})")
            
            if response.status_code == 401:
                st.error("🔑 Authentication failed. Please check your API key.")
                st.info("💡 **Fix:** Click 'Test API Connection' in sidebar or update your API key")
            elif response.status_code == 400:
                st.error("📝 Bad request. The request format may be invalid.")
                st.info("💡 **Possible causes:**\n- Prompt too long\n- Special characters in text\n- Invalid model\n\n**Try:**\n1. Test API Connection in sidebar\n2. Try a different model (e.g., openai/gpt-3.5-turbo)\n3. Shorten your resume or job description\n4. Check your API credits at openrouter.ai")
            elif response.status_code == 429:
                st.error("⏱️ Rate limit exceeded. Please wait 60 seconds and try again.")
            elif response.status_code == 500:
                st.error("🔧 Server error. The API service may be temporarily unavailable.")
                st.info("💡 **Try:** Wait a few minutes and try again, or use a different model")
            else:
                st.error(f"Error details: {error_text}")
            
            return None
        
        # Parse response
        try:
            result = response.json()
        except:
            st.error("❌ Failed to parse API response")
            return None
        
        # Validate response structure
        if 'choices' not in result:
            st.error(f"❌ Unexpected API response structure: {result}")
            return None
        
        if not result['choices']:
            st.error("❌ API returned empty choices")
            return None
        
        if 'message' not in result['choices'][0]:
            st.error("❌ API response missing message field")
            return None
        
        return result['choices'][0]['message']['content']
    
    except requests.exceptions.Timeout:
        st.error("⏱️ Request timed out (120 seconds). The API is taking too long.")
        st.info("💡 **Try:** Using a smaller resume or shorter job description")
        return None
    except requests.exceptions.ConnectionError:
        st.error("🌐 Connection error. Please check your internet connection.")
        return None
    except Exception as e:
        st.error(f"❌ Unexpected error: {str(e)}")
        return None

def clean_resume_text(text):
    """Clean and normalize resume text before PDF/DOCX generation - ENHANCED."""
    if not text:
        return ""
    
    # Remove markdown formatting
    text = text.replace('**', '').replace('##', '').replace('#', '').replace('###', '')
    
    # Remove meta-instructional text and section headers that shouldn't appear
    meta_patterns = [
        r'CRITICAL SUCCESS FACTORS:.*',
        r'FUNCTIONAL SKILLS:.*',
        r'INCLUDE ALL JD-REQUIRED.*',
        r'USE JD LANGUAGE.*',
        r'This resume is optimized.*',
        r'Achieve 85%\+ keyword match.*',
        r'Footer:.*',
        r'ATS-Optimized Resume.*Last Updated.*',
    ]
    
    for pattern in meta_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)
    
    # Fix double commas and spacing
    text = re.sub(r',,+', ',', text)
    text = re.sub(r'\s+,', ',', text)
    text = re.sub(r',\s+', ', ', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Normalize bullets
    text = text.replace('■', '•').replace('–', '-').replace('—', '|')
    
    # Remove separator lines and dashes
    text = re.sub(r'\n---+\n', '\n', text)
    text = re.sub(r'^---+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^•\s*--\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*--\s*$', '', text, flags=re.MULTILINE)
    
    # Remove lines with only special characters or very short
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip lines that are just bullets, dashes, pipes, or very short (< 3 chars)
        if stripped and len(stripped) > 3 and not re.match(r'^[•\-–—\|\*]+$', stripped):
            # Also skip lines that are all caps and look like headers but aren't section headers
            if stripped.isupper() and len(stripped) < 50:
                # Check if it's not a real section header
                section_headers = ['PROFESSIONAL SUMMARY', 'DOMAIN EXPERTISE', 'PROFESSIONAL EXPERIENCE', 
                                  'CERTIFICATIONS', 'TECHNICAL', 'KEY PROJECTS', 'EDUCATION', 'SKILLS']
                if not any(header in stripped for header in section_headers):
                    continue
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def generate_pdf_resume(resume_text, filename="ATS_Optimized_Resume.pdf"):
    """Generates ATS-friendly PDF with proper formatting."""
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_filename = temp_file.name
        temp_file.close()
        
        doc = SimpleDocTemplate(
            temp_filename, 
            pagesize=A4,
            rightMargin=0.75*inch, 
            leftMargin=0.75*inch,
            topMargin=0.75*inch, 
            bottomMargin=0.75*inch
        )
        
        styles = getSampleStyleSheet()
        
        # Styles
        style_title = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=4
        )
        
        style_subtitle = ParagraphStyle(
            'SubtitleStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=8
        )
        
        style_contact = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=10
        )
        
        style_heading = ParagraphStyle(
            'HeadingStyle',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            spaceBefore=14,
            spaceAfter=6,
            alignment=TA_LEFT,
            textColor=black,
            textTransform='uppercase'
        )
        
        style_summary = ParagraphStyle(
            'SummaryStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=13,
            alignment=TA_JUSTIFY,
            spaceAfter=10,
            leftIndent=0,
            rightIndent=0
        )
        
        style_domain = ParagraphStyle(
            'DomainStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=10,
            leftIndent=0,
            rightIndent=0
        )
        
        style_content = ParagraphStyle(
            'ContentStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
            spaceAfter=5,
            leftIndent=0
        )
        
        style_bullet = ParagraphStyle(
            'BulletStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
            spaceAfter=4,
            leftIndent=20,
            firstLineIndent=-15
        )

        # Parse and Format
        story = []
        resume_text = clean_resume_text(resume_text)
        lines = resume_text.split('\n')
        
        current_section = None
        is_first_line = True
        skip_section = False
        
        # Enhanced skip keywords - more comprehensive
        skip_keywords = [
            'critical success factors', 'achieve 85%+ keyword match',
            'include all jd-required domains', 'use jd language in professional summary',
            'this resume is optimized for ats', 'designed to achieve an 85%+',
            'include, all, jd-required, skills', 'footer:', 'footer',
            'ats-optimized resume | it business analyst | last updated',
            'last updated:', 'optimization notes', 'ai-generated',
            'meta-instructional', 'functional skills:', 'technical skills:',
            'professional skills:', 'additional information', 'references available'
        ]
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                story.append(Spacer(1, 4))
                continue
            
            line_clean = line_stripped
            line_lower = line_clean.lower()
            
            # Skip meta-instructional sections
            if any(keyword in line_lower for keyword in skip_keywords):
                skip_section = True
                continue
            
            # Reset skip flag if we hit a real section header
            if any(header in line_lower for header in ['professional summary', 'domain expertise', 
                                                         'professional experience', 'certifications',
                                                         'technical', 'key projects', 'education']):
                skip_section = False
            
            if skip_section:
                continue
            
            # Detect section headers
            section_keywords = {
                'professional summary': 'professional summary',
                'domain expertise': 'domain expertise',
                'professional experience': 'professional experience',
                'certifications': 'certifications',
                'technical': 'technical & professional skills',
                'key projects': 'key projects',
                'education': 'education'
            }
            
            is_header = False
            for keyword, section_name in section_keywords.items():
                if keyword in line_lower:
                    story.append(Spacer(1, 8))
                    story.append(Paragraph(line_clean.upper(), style_heading))
                    story.append(Spacer(1, 2))
                    current_section = section_name
                    is_header = True
                    break
            
            if is_header:
                continue
            
            # Name (first non-empty line)
            if is_first_line and len(line_clean.split()) <= 3:
                if not any(c in line_clean for c in ['@', '📧', '📞', '📍', '■', '|', '—', '•', ':', '+']):
                    story.append(Paragraph(line_clean, style_title))
                    is_first_line = False
                    continue
            
            # Title/Role
            if is_first_line and any(term in line_clean for term in ['Business Analyst', 'Agile BA', 'Requirements Engineer']):
                story.append(Paragraph(line_clean, style_subtitle))
                is_first_line = False
                continue
            
            # Contact info
            if any(icon in line_clean for icon in ['📧', '📞', '📍', '■', '@', '+91', 'linkedin', 'github', 'email']):
                if '|' in line_clean or '@' in line_clean or '+' in line_clean:
                    story.append(Paragraph(line_clean, style_contact))
                    continue
            
            # Professional Summary - Justified
            if current_section == 'professional summary':
                if not line_clean.startswith('•') and not line_clean.startswith('-'):
                    story.append(Paragraph(line_clean, style_summary))
                continue
            
            # Domain Expertise - Comma-separated with clean formatting
            if current_section == 'domain expertise':
                domains = [d.strip() for d in line_clean.replace('  ', ' ').split(',') if d.strip() and len(d.strip()) > 2]
                domain_line = ', '.join(domains)
                if domain_line and len(domain_line) > 5:
                    story.append(Paragraph(domain_line, style_domain))
                continue
            
            # Professional Experience - Left aligned with bullets
            if current_section == 'professional experience':
                if '|' in line_clean and '—' in line_clean:
                    parts = line_clean.split('|')
                    if len(parts) >= 2:
                        role = parts[0].strip()
                        rest = '|'.join(parts[1:]).strip()
                        formatted = f"<b>{role}</b> | {rest}"
                        story.append(Paragraph(formatted, style_content))
                    else:
                        story.append(Paragraph(line_clean, style_content))
                elif line_clean.startswith('•') or line_clean.startswith('-'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        bullet_text = f"• {content}"
                        story.append(Paragraph(bullet_text, style_bullet))
                else:
                    story.append(Paragraph(line_clean, style_content))
                continue
            
            # Education - Left aligned
            if current_section == 'education':
                if '—' in line_clean or '|' in line_clean:
                    story.append(Paragraph(line_clean, style_content))
                elif line_clean.startswith('•') or line_clean.startswith('-'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        story.append(Paragraph(f"• {content}", style_bullet))
                else:
                    story.append(Paragraph(line_clean, style_content))
                continue
            
            # Certifications - Left aligned with bullets
            if current_section == 'certifications':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('■') or line_clean.startswith('🏅'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        story.append(Paragraph(f"• {content}", style_bullet))
                else:
                    story.append(Paragraph(line_clean, style_content))
                continue
            
            # Technical Skills - Left aligned
            if current_section == 'technical & professional skills':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('**'):
                    content = line_clean.replace('**', '').strip()
                    if content.startswith('-') or content.startswith('•'):
                        content = content[1:].strip()
                    if content and len(content) > 5:
                        story.append(Paragraph(f"• {content}", style_bullet))
                else:
                    story.append(Paragraph(line_clean, style_content))
                continue
            
            # Key Projects - Left aligned with bullets
            if current_section == 'key projects':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('■') or line_clean.startswith('📌'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        story.append(Paragraph(f"• {content}", style_bullet))
                else:
                    story.append(Paragraph(line_clean, style_content))
                continue
            
            # Default fallback
            story.append(Paragraph(line_clean, style_content))
        
        doc.build(story)
        return temp_filename
        
    except Exception as e:
        st.error(f"❌ Error generating PDF: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

def generate_word_resume(resume_text, filename="ATS_Optimized_Resume.docx"):
    """Generates ATS-friendly Word document with proper formatting."""
    try:
        doc = Document()
        
        # Set default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        
        # Parse and Format
        resume_text = clean_resume_text(resume_text)
        lines = resume_text.split('\n')
        
        current_section = None
        is_first_line = True
        skip_section = False
        
        # Enhanced skip keywords
        skip_keywords = [
            'critical success factors', 'achieve 85%+ keyword match',
            'include all jd-required domains', 'use jd language in professional summary',
            'this resume is optimized for ats', 'designed to achieve an 85%+',
            'include, all, jd-required, skills', 'footer:', 'footer',
            'ats-optimized resume | it business analyst | last updated',
            'last updated:', 'optimization notes', 'ai-generated',
            'meta-instructional', 'functional skills:', 'technical skills:',
            'professional skills:', 'additional information', 'references available'
        ]
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                doc.add_paragraph()
                continue
            
            line_clean = line_stripped
            line_lower = line_clean.lower()
            
            # Skip meta-instructional sections
            if any(keyword in line_lower for keyword in skip_keywords):
                skip_section = True
                continue
            
            # Reset skip flag if we hit a real section header
            if any(header in line_lower for header in ['professional summary', 'domain expertise', 
                                                         'professional experience', 'certifications',
                                                         'technical', 'key projects', 'education']):
                skip_section = False
            
            if skip_section:
                continue
            
            # Detect section headers
            section_keywords = {
                'professional summary': 'professional summary',
                'domain expertise': 'domain expertise',
                'professional experience': 'professional experience',
                'certifications': 'certifications',
                'technical': 'technical & professional skills',
                'key projects': 'key projects',
                'education': 'education'
            }
            
            is_header = False
            for keyword, section_name in section_keywords.items():
                if keyword in line_lower:
                    # Add section header
                    p = doc.add_paragraph(line_clean.upper())
                    p.style = 'Heading 2'
                    is_header = True
                    current_section = section_name
                    break
            
            if is_header:
                continue
            
            # Name (first non-empty line)
            if is_first_line and len(line_clean.split()) <= 3:
                if not any(c in line_clean for c in ['@', '📧', '📞', '📍', '■', '|', '—', '•', ':', '+']):
                    p = doc.add_paragraph(line_clean)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style.font.size = Pt(16)
                    p.style.font.bold = True
                    is_first_line = False
                    continue
            
            # Title/Role
            if is_first_line and any(term in line_clean for term in ['Business Analyst', 'Agile BA', 'Requirements Engineer']):
                p = doc.add_paragraph(line_clean)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style.font.size = Pt(11)
                is_first_line = False
                continue
            
            # Contact info
            if any(icon in line_clean for icon in ['📧', '📞', '📍', '■', '@', '+91', 'linkedin', 'github', 'email']):
                if '|' in line_clean or '@' in line_clean or '+' in line_clean:
                    p = doc.add_paragraph(line_clean)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style.font.size = Pt(9)
                    continue
            
            # Professional Summary - Justified
            if current_section == 'professional summary':
                if not line_clean.startswith('•') and not line_clean.startswith('-'):
                    p = doc.add_paragraph(line_clean)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                continue
            
            # Domain Expertise - Comma-separated
            if current_section == 'domain expertise':
                domains = [d.strip() for d in line_clean.replace('  ', ' ').split(',') if d.strip() and len(d.strip()) > 2]
                domain_line = ', '.join(domains)
                if domain_line and len(domain_line) > 5:
                    doc.add_paragraph(domain_line)
                continue
            
            # Professional Experience - Left aligned with bullets
            if current_section == 'professional experience':
                if '|' in line_clean and '—' in line_clean:
                    parts = line_clean.split('|')
                    if len(parts) >= 2:
                        role = parts[0].strip()
                        rest = '|'.join(parts[1:]).strip()
                        p = doc.add_paragraph()
                        runner = p.add_run(f"{role} | {rest}")
                        runner.bold = True
                    else:
                        doc.add_paragraph(line_clean)
                elif line_clean.startswith('•') or line_clean.startswith('-'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        p = doc.add_paragraph(content, style='List Bullet')
                else:
                    doc.add_paragraph(line_clean)
                continue
            
            # Education - Left aligned
            if current_section == 'education':
                if '—' in line_clean or '|' in line_clean:
                    doc.add_paragraph(line_clean)
                elif line_clean.startswith('•') or line_clean.startswith('-'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        p = doc.add_paragraph(content, style='List Bullet')
                else:
                    doc.add_paragraph(line_clean)
                continue
            
            # Certifications - Left aligned with bullets
            if current_section == 'certifications':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('■') or line_clean.startswith('🏅'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        p = doc.add_paragraph(content, style='List Bullet')
                else:
                    doc.add_paragraph(line_clean)
                continue
            
            # Technical Skills - Left aligned
            if current_section == 'technical & professional skills':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('**'):
                    content = line_clean.replace('**', '').strip()
                    if content.startswith('-') or content.startswith('•'):
                        content = content[1:].strip()
                    if content and len(content) > 5:
                        p = doc.add_paragraph(content, style='List Bullet')
                else:
                    doc.add_paragraph(line_clean)
                continue
            
            # Key Projects - Left aligned with bullets
            if current_section == 'key projects':
                if line_clean.startswith('•') or line_clean.startswith('-') or line_clean.startswith('■') or line_clean.startswith('📌'):
                    content = line_clean[1:].strip()
                    if content and len(content) > 10:
                        p = doc.add_paragraph(content, style='List Bullet')
                else:
                    doc.add_paragraph(line_clean)
                continue
            
            # Default fallback
            doc.add_paragraph(line_clean)
        
        # Save to temp file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_filename = temp_file.name
        temp_file.close()
        
        doc.save(temp_filename)
        return temp_filename
        
    except Exception as e:
        st.error(f"❌ Error generating Word document: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
        return None

# ============================================
# Prompts - Optimized for API
# ============================================

VALIDATION_SYSTEM_PROMPT = """You are an expert IT Business Analyst Hiring Manager. Compare the candidate's resume against the job description and output a validation report in the exact ASCII table format shown."""

VALIDATION_USER_PROMPT = """
Resume:
{resume_text}

Job Description:
{jd_text}

Generate validation report in this EXACT format:

╔══════════════════════════════════════════════════════════════════════════════╗
║                 IT-BUSINESS ANALYST RESUME VALIDATION REPORT                  ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║   Candidate: [Name]                                                          ║
║   Job Position: [Position]                                                   ║
║   Analysis Date: [Date]                                                      ║
║                                                                              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║   🎯 OVERALL ELIGIBILITY: [XX]%                                              ║
║   ✅ GOOD MATCH - RECOMMENDED                                                 ║
║                                                                              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║   📈 SCORE BREAKDOWN:                                                         ║
║                                                                              ║
║   • Education:           [XX]%                                               ║
║   • Functional Skills:   [XX]%                                               ║
║   • Experience:          [XX]%                                               ║
║                                                                              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║   🔍 DETAILED ANALYSIS:                                                       ║
║                                                                              ║
║   ✅ Matched Skills:                                                          ║
║      • [List skills]                                                          ║
║                                                                              ║
║   ❌ Missing Skills:                                                          ║
║      • [List skills]                                                          ║
║                                                                              ║
║   ✅ Matched Experience:                                                      ║
║      • [List experience]                                                      ║
║                                                                              ║
║   ❌ Missing Experience:                                                      ║
║      • [List gaps]                                                            ║
║                                                                              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║   💡 RECOMMENDATIONS:                                                         ║
║   1. [Recommendation 1]                                                      ║
║   2. [Recommendation 2]                                                      ║
║   3. [Recommendation 3]                                                      ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

RESUME_GEN_SYSTEM_PROMPT = """You are an expert Resume Writer. Create an ATS-optimized resume that achieves 85%+ match with the job description. Follow the Priya Sharma template format exactly. NO footers, NO meta-text, NO all-caps sections, NO instructional text."""

RESUME_GEN_USER_PROMPT = """
Original Resume:
{resume_text}

Job Description:
{jd_text}

Create ATS-optimized resume following this structure:
1. Name
2. Title (IT Business Analyst | Agile BA | Requirements Engineer)
3. Contact Info
4. Professional Summary (4-5 lines)
5. Domain Expertise (comma-separated)
6. Professional Experience (Role | Company — Location | Date | Bullets)
7. Certifications
8. Technical & Professional Skills
9. Key Projects
10. Education

CRITICAL RULES:
- Combine candidate's domains with job-required domains
- Include 90%+ of JD keywords
- Use standard formatting, NO markdown
- NO double commas
- NO footers
- NO meta-text like "Critical Success Factors" or "Functional Skills:"
- Output ONLY the resume content - nothing else
"""

# ============================================
# Main Application
# ============================================

def main():
    # Title and Header
    st.title("📊 Business Analyst Job Apply Pro")
    st.markdown("### ATS Optimized Resume Validator & Generator")
    st.markdown("---")
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key
        api_key = get_api_key()
        
        if not api_key:
            st.warning("⚠️ Please enter your OpenRouter API Key")
            st.markdown("[Get your API key here](https://openrouter.ai/keys)")
            st.stop()
        else:
            # Test API connection
            if st.button("🔍 Test API Connection", key=f"test_api_{st.session_state.reset_counter}"):
                test_api_connection(api_key)
            st.success("✅ API Key configured")
        
        # Model Selection
        st.markdown("---")
        st.markdown("### 🤖 Model Settings")
        model = st.selectbox(
            "AI Model",
            ["qwen/qwen-2.5-72b-instruct", "qwen/qwen-2.5-coder-32b-instruct", "openai/gpt-3.5-turbo", "anthropic/claude-3-haiku"],
            index=0,
            help="Qwen model for best results",
            key=f"model_select_{st.session_state.reset_counter}"
        )
        
        # About Section
        st.markdown("---")
        st.markdown("### ℹ️ About")
        st.markdown("""
        **Version:** 1.0.0
        
        **Powered by:**
        - Streamlit
        - OpenRouter API
        - Qwen Models
        """)
    
    # Main Content - Two Columns
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("1️⃣ Upload Your Resume")
        uploaded_file = st.file_uploader(
            "Choose a file", 
            type=['pdf', 'docx', 'txt'],
            help="Upload your current resume (PDF, DOCX, or TXT)",
            key=f"resume_uploader_{st.session_state.reset_counter}"
        )
        
        if uploaded_file:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            st.info(f"📄 Size: {uploaded_file.size / 1024:.1f} KB")
        
        st.markdown("---")
        st.subheader("2️⃣ Job Requirements")
        jd_text = st.text_area(
            "Paste Employer's Job Description", 
            height=250, 
            placeholder="Paste the full job description...",
            help="Copy and paste the complete job posting",
            key=f"jd_textarea_{st.session_state.reset_counter}"
        )
        
        if jd_text:
            st.success(f"✅ Job description loaded ({len(jd_text)} characters)")
    
    with col2:
        st.subheader("3️⃣ Actions")
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            validate_btn = st.button("🔍 Validate Resume", type="primary", use_container_width=True, 
                                   key=f"validate_btn_{st.session_state.reset_counter}")
        with col_btn2:
            reset_btn = st.button("🔄 Reset All", use_container_width=True, key="reset_button")
        
        # Reset functionality
        if reset_btn:
            st.session_state.reset_counter += 1
            st.session_state.validation_report = None
            st.session_state.generated_resume = None
            st.session_state.resume_text = None
            st.session_state.processing = False
            st.success("✅ All data cleared! Ready for fresh operation.")
            st.rerun()
        
        st.markdown("---")
        st.subheader("📊 Status")
    
    # Main Processing Logic
    if uploaded_file and jd_text:
        # Extract resume text
        if st.session_state.resume_text is None:
            with st.spinner("📄 Parsing resume..."):
                st.session_state.resume_text = extract_text_from_file(uploaded_file)
                if st.session_state.resume_text:
                    st.success("✅ Resume parsed successfully!")
        
        # Validation Button
        if validate_btn:
            if not st.session_state.resume_text:
                st.error("❌ Could not parse resume. Please try uploading again.")
            else:
                with st.spinner("🔍 Analyzing resume against job description..."):
                    st.session_state.processing = True
                    
                    # Prepare prompt with size limits
                    resume_text = st.session_state.resume_text[:2500]  # Limit to 2500 chars
                    jd_text_limited = jd_text[:1500]  # Limit to 1500 chars
                    
                    prompt = VALIDATION_USER_PROMPT.format(
                        resume_text=resume_text,
                        jd_text=jd_text_limited
                    )
                    
                    # Call API
                    report = call_openrouter_api(
                        prompt, 
                        VALIDATION_SYSTEM_PROMPT, 
                        api_key,
                        model,
                        max_tokens=2500
                    )
                    
                    st.session_state.processing = False
                    
                    if report:
                        st.session_state.validation_report = report
                        st.success("✅ Validation complete!")
                        st.rerun()
                    else:
                        st.error("❌ Failed to generate validation report.")
                        st.info("💡 **Troubleshooting:**\n1. Click 'Test API Connection' in sidebar\n2. Check your API credits at openrouter.ai\n3. Try a different model (e.g., openai/gpt-3.5-turbo)\n4. Shorten your resume or job description")
        
        # Display Validation Report
        if st.session_state.validation_report:
            st.markdown("---")
            st.subheader("📋 Validation Report")
            st.code(st.session_state.validation_report, language="text")
            
            st.markdown("---")
            st.subheader("4️⃣ Update Resume?")
            
            update_choice = st.radio(
                "Do you want to update your resume?", 
                ["Yes, generate optimized resume (85%+ match)", "No, I'm good with current resume"], 
                horizontal=True,
                key=f"update_choice_{st.session_state.reset_counter}"
            )
            
            if update_choice == "Yes, generate optimized resume (85%+ match)":
                if st.button("✨ Generate ATS Optimized Resume", type="primary",
                           key=f"generate_btn_{st.session_state.reset_counter}"):
                    with st.spinner("✍️ Generating new ATS-optimized resume..."):
                        # Prepare prompt with size limits
                        resume_text = st.session_state.resume_text[:3500]
                        jd_text_limited = jd_text[:2500]
                        
                        prompt = RESUME_GEN_USER_PROMPT.format(
                            resume_text=resume_text,
                            jd_text=jd_text_limited
                        )
                        
                        new_resume = call_openrouter_api(
                            prompt, 
                            RESUME_GEN_SYSTEM_PROMPT, 
                            api_key,
                            model,
                            max_tokens=4000
                        )
                        
                        if new_resume:
                            st.session_state.generated_resume = new_resume
                            st.success("✅ Resume generated successfully!")
                            st.rerun()
                        else:
                            st.error("❌ Failed to generate resume.")
                            st.info("💡 **Troubleshooting:**\n1. Check your API credits at openrouter.ai\n2. Try a different model (e.g., openai/gpt-3.5-turbo)\n3. Shorten the inputs\n4. Wait 60 seconds if rate limited")
    
    elif uploaded_file and not jd_text:
        st.info("📝 **Please paste the Job Description to proceed.**")
    elif jd_text and not uploaded_file:
        st.info("📄 **Please upload your Resume to proceed.**")
    elif not uploaded_file and not jd_text:
        st.info("👆 **Upload your resume and paste the job description to get started.**")
    
    # Display Generated Resume
    if st.session_state.generated_resume:
        st.markdown("---")
        st.subheader("📄 Generated ATS Resume")
        
        with st.expander("👁️ Preview Resume Text", expanded=True):
            st.text_area("Resume Preview", value=st.session_state.generated_resume, height=500,
                       key=f"preview_text_{st.session_state.reset_counter}")
        
        # Generate PDF and Word
        col_pdf, col_word = st.columns(2)
        
        with col_pdf:
            with st.spinner("📄 Creating PDF..."):
                pdf_file = generate_pdf_resume(st.session_state.generated_resume, "ATS_Optimized_Resume.pdf")
                
                if pdf_file and os.path.exists(pdf_file):
                    st.success("✅ PDF generated!")
                    
                    with open(pdf_file, "rb") as f:
                        pdf_bytes = f.read()
                    
                    st.download_button(
                        label="📥 Download PDF",
                        data=pdf_bytes,
                        file_name="ATS_Optimized_Resume.pdf",
                        mime="application/pdf",
                        type="primary",
                        key=f"download_pdf_{st.session_state.reset_counter}"
                    )
                    
                    try:
                        os.remove(pdf_file)
                    except:
                        pass
        
        with col_word:
            with st.spinner("📝 Creating Word..."):
                word_file = generate_word_resume(st.session_state.generated_resume, "ATS_Optimized_Resume.docx")
                
                if word_file and os.path.exists(word_file):
                    st.success("✅ Word generated!")
                    
                    with open(word_file, "rb") as f:
                        word_bytes = f.read()
                    
                    st.download_button(
                        label="📥 Download Word",
                        data=word_bytes,
                        file_name="ATS_Optimized_Resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        key=f"download_word_{st.session_state.reset_counter}"
                    )
                    
                    try:
                        os.remove(word_file)
                    except:
                        pass
        
        st.markdown("---")
        st.success("🎉 **Your ATS-optimized resume is ready! Download in PDF or Word format.**")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 12px;">
        <p>Built with ❤️ for Business Analysts worldwide | Powered by Streamlit & OpenRouter</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"❌ Application error: {str(e)}")
        st.exception(e)
 
